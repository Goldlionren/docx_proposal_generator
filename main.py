import os
import shutil
import subprocess
import sys
import tempfile
import tkinter as tk
from copy import deepcopy
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext, ttk
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


APP_NAME = "DOCX Proposal Generator"
APP_DIR = Path(__file__).resolve().parent
DOCX_FORMAT = 16
PAGE_BREAK = 7


class ProposalGenerationError(Exception):
    pass


def find_pandoc():
    candidate_paths = [
        APP_DIR / ".venv" / "Scripts" / "pandoc.exe",
        Path(sys.prefix) / "Scripts" / "pandoc.exe",
        APP_DIR / "pandoc.exe",
        APP_DIR / "tools" / "pandoc.exe",
        APP_DIR / "tools" / "pandoc" / "pandoc.exe",
    ]

    for candidate_path in candidate_paths:
        if candidate_path.is_file():
            return str(candidate_path)

    try:
        import pypandoc

        pandoc_path = pypandoc.get_pandoc_path()
        if pandoc_path:
            pypandoc_candidate = Path(pandoc_path)
            if pypandoc_candidate.is_file():
                return str(pypandoc_candidate)

            exe_candidate = pypandoc_candidate.with_suffix(".exe")
            if exe_candidate.is_file():
                return str(exe_candidate)
    except (ImportError, OSError):
        pass

    return shutil.which("pandoc")


def check_pandoc_version(pandoc_path):
    try:
        result = subprocess.run(
            [pandoc_path, "--version"],
            capture_output=True,
            text=True,
            check=False,
        )
    except OSError as exc:
        raise ProposalGenerationError(f"Failed to run Pandoc: {exc}") from exc

    if result.returncode != 0:
        raise ProposalGenerationError(
            "Pandoc is not installed in the local Python environment.\n"
            "Please run: .\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt"
        )

    return result.stdout.splitlines()[0] if result.stdout else "pandoc"


def convert_markdown_to_docx(pandoc_path, source_md_path, reference_docx_path, body_docx_path):
    pandoc_command = [
        pandoc_path,
        source_md_path,
        "--reference-doc",
        reference_docx_path,
        "-o",
        body_docx_path,
    ]
    result = subprocess.run(
        pandoc_command,
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        raise ProposalGenerationError(
            "Pandoc conversion failed.\n\n"
            f"Return code: {result.returncode}\n\n"
            f"{result.stderr.strip() or 'No error output was returned.'}"
        )
    return pandoc_command, result


def compose_with_word(front_template_path, body_docx_path, output_docx_path, metadata, log):
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise ProposalGenerationError("pywin32 is not installed.") from exc

    word = None
    doc = None
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        template_path = Path(front_template_path).resolve()
        if template_path.suffix.lower() == ".dotx":
            doc = word.Documents.Add(Template=str(template_path))
            log(f"Created new document from DOTX template: {template_path}")
        else:
            doc = word.Documents.Open(str(template_path))
            log(f"Opened DOCX front matter template: {template_path}")

        replace_placeholders_everywhere(doc, metadata, log)

        content_range = find_word_placeholder(doc, "{{BODY_CONTENT}}")
        if content_range is None:
            raise ProposalGenerationError(
                "The front matter template does not contain {{BODY_CONTENT}}."
            )

        content_range.Text = ""
        insertion_range = content_range.Duplicate
        insertion_range.Collapse(0)
        insertion_range.InsertFile(str(Path(body_docx_path).resolve()))

        update_toc_without_creating_duplicate(doc, log)
        update_word_fields(doc, log)
        doc.SaveAs2(str(Path(output_docx_path).resolve()), FileFormat=DOCX_FORMAT)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def replace_placeholders_everywhere(doc, replacements, log=None):
    safe_replacements = {
        placeholder: value
        for placeholder, value in replacements.items()
        if placeholder != "{{BODY_CONTENT}}"
    }
    if not safe_replacements:
        return

    replaced_count = 0
    replaced_count += replace_placeholders_in_word_range(doc.Content, safe_replacements)

    for story_range in doc.StoryRanges:
        try:
            current_range = story_range
            while current_range is not None:
                replaced_count += replace_placeholders_in_word_range(
                    current_range,
                    safe_replacements,
                )
                current_range = current_range.NextStoryRange
        except Exception as exc:
            if log:
                log(f"Skipped one Word story while replacing placeholders: {exc}")

    replaced_count += replace_placeholders_in_shapes(getattr(doc, "Shapes", None), safe_replacements, log)

    for section in doc.Sections:
        for header in section.Headers:
            replaced_count += replace_placeholders_in_word_range(header.Range, safe_replacements)
            replaced_count += replace_placeholders_in_shapes(getattr(header, "Shapes", None), safe_replacements, log)
        for footer in section.Footers:
            replaced_count += replace_placeholders_in_word_range(footer.Range, safe_replacements)
            replaced_count += replace_placeholders_in_shapes(getattr(footer, "Shapes", None), safe_replacements, log)

    if log:
        log(f"Placeholder replacement complete. Replacements attempted: {replaced_count}.")


def replace_placeholders_in_word_range(word_range, replacements):
    replacement_count = 0
    for placeholder, value in replacements.items():
        if replace_placeholder_in_word_range(word_range, placeholder, value):
            replacement_count += 1
    return replacement_count


def replace_placeholder_in_word_range(word_range, placeholder, value):
    find = word_range.Find
    find.ClearFormatting()
    find.Replacement.ClearFormatting()
    return bool(find.Execute(
        FindText=placeholder,
        ReplaceWith=value,
        Replace=2,
        Forward=True,
        MatchCase=False,
        MatchWholeWord=False,
        MatchWildcards=False,
        Wrap=1,
    ))


def replace_placeholders_in_shapes(shapes, replacements, log=None):
    if shapes is None:
        return 0

    replacement_count = 0
    try:
        shape_count = shapes.Count
    except Exception:
        return 0

    for index in range(1, shape_count + 1):
        try:
            shape = shapes.Item(index)
            replacement_count += replace_placeholders_in_shape(shape, replacements, log)
        except Exception as exc:
            if log:
                log(f"Skipped one shape while replacing placeholders: {exc}")

    return replacement_count


def replace_placeholders_in_shape(shape, replacements, log=None):
    replacement_count = 0

    try:
        text_frame = shape.TextFrame
        if text_frame.HasText:
            replacement_count += replace_placeholders_in_word_range(
                text_frame.TextRange,
                replacements,
            )
    except Exception:
        pass

    try:
        text_frame2 = shape.TextFrame2
        if text_frame2.HasText:
            replacement_count += replace_placeholders_in_text_frame2(
                text_frame2,
                replacements,
            )
    except Exception:
        pass

    for child_collection_name in ("GroupItems", "CanvasItems"):
        try:
            child_shapes = getattr(shape, child_collection_name)
            replacement_count += replace_placeholders_in_shapes(
                child_shapes,
                replacements,
                log,
            )
        except Exception:
            pass

    return replacement_count


def replace_placeholders_in_text_frame2(text_frame2, replacements):
    replacement_count = 0
    try:
        text_range = text_frame2.TextRange
        original_text = text_range.Text
    except Exception:
        return 0

    new_text = original_text
    for placeholder, value in replacements.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, value)
            replacement_count += 1

    if new_text != original_text:
        text_range.Text = new_text

    return replacement_count


def find_word_placeholder(doc, placeholder):
    search_range = doc.Content
    find = search_range.Find
    find.ClearFormatting()
    found = find.Execute(
        FindText=placeholder,
        Forward=True,
        MatchCase=False,
        MatchWholeWord=False,
        MatchWildcards=False,
        Wrap=0,
    )
    return search_range if found else None


def update_toc_without_creating_duplicate(doc, log=None):
    toc_count = doc.TablesOfContents.Count
    if toc_count == 0:
        if log:
            log("No existing Word TOC found in the template. No duplicate TOC was created.")
        return

    for index in range(toc_count, 1, -1):
        try:
            doc.TablesOfContents(index).Delete()
            if log:
                log(f"Deleted duplicate TOC at index {index}.")
        except Exception:
            try:
                doc.TablesOfContents(index).Range.Delete()
                if log:
                    log(f"Deleted duplicate TOC range at index {index}.")
            except Exception as exc:
                if log:
                    log(f"Could not delete duplicate TOC at index {index}: {exc}")

    doc.TablesOfContents(1).Update()
    if log:
        log("Updated existing template TOC.")


def update_word_fields(doc, log):
    doc.Fields.Update()

    for story_range in doc.StoryRanges:
        try:
            story_range.Fields.Update()
        except Exception:
            pass

    for section in doc.Sections:
        for header in section.Headers:
            header.Range.Fields.Update()
        for footer in section.Footers:
            footer.Range.Fields.Update()

    log("Microsoft Word updated fields, headers, footers, and page numbers.")


def compose_without_word(front_template_path, body_docx_path, output_docx_path, metadata, log):
    output_path = Path(output_docx_path)
    temp_output = output_path.with_suffix(".fallback.tmp.docx")
    converted_template_path = None
    try:
        template_path = Path(front_template_path)
        if template_path.suffix.lower() == ".dotx":
            converted_template_path = output_path.with_suffix(".front-template.tmp.docx")
            convert_dotx_to_docx_package(template_path, converted_template_path)
            doc = Document(converted_template_path)
            log(f"Loaded DOTX front matter template through temporary DOCX copy: {template_path}")
        else:
            doc = Document(front_template_path)

        body_doc = Document(body_docx_path)

        replace_placeholders_python_docx(doc, metadata)
        marker_paragraph = find_body_marker_paragraph(doc)
        if marker_paragraph is None:
            raise ProposalGenerationError(
                "The front matter template does not contain {{BODY_CONTENT}}."
            )

        clear_paragraph_text(marker_paragraph)
        if document_has_toc_field(doc):
            log("Existing TOC field found in the front matter template. No duplicate TOC was created.")
        else:
            insert_toc_before_paragraph(marker_paragraph)
            log("No existing TOC field found. Added fallback TOC field before the body.")
        insert_body_after_paragraph(marker_paragraph, body_doc)
        marker_paragraph._element.getparent().remove(marker_paragraph._element)

        doc.save(temp_output)
        replace_placeholders_in_docx_package(temp_output, metadata)
        merge_docx_media_relationships(temp_output, body_docx_path, output_path)
        log(
            "Microsoft Word was not available. Generated DOCX with a TOC field, "
            "but the TOC and page numbers may need to be updated manually in Word."
        )
    finally:
        if temp_output.exists():
            try:
                temp_output.unlink()
            except OSError:
                pass
        if converted_template_path and converted_template_path.exists():
            try:
                converted_template_path.unlink()
            except OSError:
                pass


def convert_dotx_to_docx_package(dotx_path, docx_path):
    with ZipFile(dotx_path, "r") as source_zip:
        entries = {name: source_zip.read(name) for name in source_zip.namelist()}

    content_types = ET.fromstring(entries["[Content_Types].xml"])
    namespace = {"ct": "http://schemas.openxmlformats.org/package/2006/content-types"}
    template_type = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.template.main+xml"
    )
    document_type = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document.main+xml"
    )

    for override in content_types.findall("ct:Override", namespace):
        if (
            override.attrib.get("PartName") == "/word/document.xml"
            and override.attrib.get("ContentType") == template_type
        ):
            override.set("ContentType", document_type)

    entries["[Content_Types].xml"] = ET.tostring(
        content_types,
        encoding="utf-8",
        xml_declaration=True,
    )

    with ZipFile(docx_path, "w", ZIP_DEFLATED) as output_zip:
        for name, data in entries.items():
            output_zip.writestr(name, data)


def replace_placeholders_in_docx_package(docx_path, replacements):
    safe_replacements = {
        placeholder: value
        for placeholder, value in replacements.items()
        if placeholder != "{{BODY_CONTENT}}"
    }
    if not safe_replacements:
        return

    path = Path(docx_path)
    temp_path = path.with_suffix(".replace.tmp.docx")
    with ZipFile(path, "r") as source_zip:
        entries = {name: source_zip.read(name) for name in source_zip.namelist()}

    for name, data in list(entries.items()):
        if not (name.startswith("word/") and name.endswith(".xml")):
            continue

        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            continue

        new_text = text
        for placeholder, value in safe_replacements.items():
            new_text = new_text.replace(placeholder, escape_xml_text(value))

        if new_text != text:
            entries[name] = new_text.encode("utf-8")

    with ZipFile(temp_path, "w", ZIP_DEFLATED) as output_zip:
        for name, data in entries.items():
            output_zip.writestr(name, data)

    temp_path.replace(path)


def escape_xml_text(value):
    return (
        str(value)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def replace_placeholders_python_docx(doc, metadata):
    for paragraph in iter_all_paragraphs(doc):
        replace_in_paragraph(paragraph, metadata)


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in doc.sections:
        for part in [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]:
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def replace_in_paragraph(paragraph, metadata):
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(placeholder in full_text for placeholder in metadata):
        return

    new_text = full_text
    for placeholder, value in metadata.items():
        new_text = new_text.replace(placeholder, value)

    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def find_body_marker_paragraph(doc):
    for paragraph in iter_all_paragraphs(doc):
        if "{{BODY_CONTENT}}" in paragraph.text:
            return paragraph
    return None


def clear_paragraph_text(paragraph):
    for run in paragraph.runs:
        run.text = ""


def document_has_toc_field(doc):
    return "TOC \\" in doc.element.xml or "TOC  \\" in doc.element.xml


def insert_toc_before_paragraph(paragraph):
    heading = OxmlElement("w:p")
    heading_ppr = OxmlElement("w:pPr")
    heading_style = OxmlElement("w:pStyle")
    heading_style.set(qn("w:val"), "Heading1")
    heading_ppr.append(heading_style)
    heading.append(heading_ppr)

    heading_run = OxmlElement("w:r")
    heading_text = OxmlElement("w:t")
    heading_text.text = "Table of Contents"
    heading_run.append(heading_text)
    heading.append(heading_run)

    toc_paragraph = OxmlElement("w:p")
    run_begin = OxmlElement("w:r")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run_begin.append(field_begin)

    run_instr = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    run_instr.append(instr_text)

    run_sep = OxmlElement("w:r")
    field_sep = OxmlElement("w:fldChar")
    field_sep.set(qn("w:fldCharType"), "separate")
    run_sep.append(field_sep)

    run_placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Right-click and update field to refresh this table of contents."
    run_placeholder.append(placeholder_text)

    run_end = OxmlElement("w:r")
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run_end.append(field_end)

    for run in [run_begin, run_instr, run_sep, run_placeholder, run_end]:
        toc_paragraph.append(run)

    page_break = OxmlElement("w:p")
    break_run = OxmlElement("w:r")
    break_element = OxmlElement("w:br")
    break_element.set(qn("w:type"), "page")
    break_run.append(break_element)
    page_break.append(break_run)

    paragraph._element.addprevious(heading)
    paragraph._element.addprevious(toc_paragraph)
    paragraph._element.addprevious(page_break)


def insert_body_after_paragraph(paragraph, body_doc):
    target = paragraph._element
    for element in reversed(list(body_doc.element.body)):
        if element.tag.endswith("}sectPr"):
            continue
        target.addnext(deepcopy(element))


def merge_docx_media_relationships(fallback_docx_path, body_docx_path, output_docx_path):
    body_media = {}
    with ZipFile(body_docx_path, "r") as body_zip:
        for name in body_zip.namelist():
            if name.startswith("word/media/"):
                body_media[name] = body_zip.read(name)

    with ZipFile(fallback_docx_path, "r") as fallback_zip:
        entries = {name: fallback_zip.read(name) for name in fallback_zip.namelist()}

    entries.update({name: data for name, data in body_media.items() if name not in entries})

    with ZipFile(output_docx_path, "w", ZIP_DEFLATED) as output_zip:
        for name, data in entries.items():
            output_zip.writestr(name, data)


class DOCXProposalGenerator(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title(APP_NAME)
        self.geometry("940x680")
        self.minsize(820, 560)

        self.front_template_docx = tk.StringVar()
        self.reference_docx = tk.StringVar()
        self.source_markdown = tk.StringVar()
        self.output_docx = tk.StringVar()
        self.client_name = tk.StringVar()
        self.document_name = tk.StringVar()
        self.author = tk.StringVar()

        self._build_ui()

    def _build_ui(self):
        root = ttk.Frame(self, padding=16)
        root.pack(fill=tk.BOTH, expand=True)

        title = ttk.Label(root, text=APP_NAME, font=("Segoe UI", 16, "bold"))
        title.grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 16))

        root.columnconfigure(1, weight=1)
        root.rowconfigure(9, weight=1)

        self._add_file_row(
            root,
            row=1,
            label="Front Matter Template DOCX/DOTX:",
            variable=self.front_template_docx,
            button_text="Browse Front Matter",
            command=self.browse_front_template_docx,
        )
        self._add_file_row(
            root,
            row=2,
            label="Reference DOCX:",
            variable=self.reference_docx,
            button_text="Browse Reference DOCX",
            command=self.browse_reference_docx,
        )
        self._add_file_row(
            root,
            row=3,
            label="Source Markdown:",
            variable=self.source_markdown,
            button_text="Browse Source MD",
            command=self.browse_source_markdown,
        )
        self._add_file_row(
            root,
            row=4,
            label="Output DOCX:",
            variable=self.output_docx,
            button_text="Save Output As",
            command=self.browse_output_docx,
        )

        self._add_text_row(root, 5, "Client Name:", self.client_name)
        self._add_text_row(root, 6, "Document Name:", self.document_name)
        self._add_text_row(root, 7, "Author:", self.author)

        start_button = ttk.Button(root, text="Start", command=self.start_conversion)
        start_button.grid(row=8, column=0, sticky="w", pady=(12, 12))

        status_label = ttk.Label(root, text="Status:")
        status_label.grid(row=9, column=0, sticky="nw", pady=(0, 4))

        self.status_text = scrolledtext.ScrolledText(root, height=14, wrap=tk.WORD)
        self.status_text.grid(row=9, column=1, columnspan=2, sticky="nsew")
        self.status_text.configure(state=tk.DISABLED)

    def _add_file_row(self, parent, row, label, variable, button_text, command):
        ttk.Label(parent, text=label).grid(row=row, column=0, sticky="w", pady=6)

        entry = ttk.Entry(parent, textvariable=variable, state="readonly")
        entry.grid(row=row, column=1, sticky="ew", padx=(8, 8), pady=6)

        button = ttk.Button(parent, text=button_text, command=command)
        button.grid(row=row, column=2, sticky="ew", pady=6)

    def _add_text_row(self, parent, row, label, variable):
        ttk.Label(parent, text=label).grid(row=row, column=0, sticky="w", pady=6)
        entry = ttk.Entry(parent, textvariable=variable)
        entry.grid(row=row, column=1, columnspan=2, sticky="ew", padx=(8, 0), pady=6)

    def browse_front_template_docx(self):
        path = filedialog.askopenfilename(
            title="Select Front Matter Template",
            filetypes=[
                ("Word Templates/Documents", "*.docx *.dotx"),
                ("Word Documents", "*.docx"),
                ("Word Templates", "*.dotx"),
            ],
        )
        if path:
            self.front_template_docx.set(path)
            self.log(f"Selected front matter template: {path}")

    def browse_reference_docx(self):
        path = filedialog.askopenfilename(
            title="Select Reference DOCX",
            filetypes=[("Word Documents", "*.docx")],
        )
        if path:
            self.reference_docx.set(path)
            self.log(f"Selected reference file: {path}")

    def browse_source_markdown(self):
        path = filedialog.askopenfilename(
            title="Select Source Markdown",
            filetypes=[
                ("Markdown Files", "*.md *.markdown"),
                ("All Files", "*.*"),
            ],
        )
        if path:
            self.source_markdown.set(path)
            self.log(f"Selected source file: {path}")

    def browse_output_docx(self):
        path = filedialog.asksaveasfilename(
            title="Save Output DOCX As",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
        )
        if path:
            self.output_docx.set(path)
            self.log(f"Selected output file: {path}")

    def start_conversion(self):
        front_template_docx = self.front_template_docx.get().strip()
        reference_docx = self.reference_docx.get().strip()
        source_markdown = self.source_markdown.get().strip()
        output_docx = self.output_docx.get().strip()

        if not self._is_file(front_template_docx):
            self.show_error("Please select a valid front matter template file.")
            return

        if not self._has_front_template_extension(front_template_docx):
            self.show_error("Please select a front matter template ending in .docx or .dotx.")
            return

        if not self._is_file(reference_docx):
            self.show_error("Please select a valid reference DOCX file.")
            return

        if not self._has_docx_extension(reference_docx):
            self.show_error("The Pandoc reference document must be a .docx file, not .dotx.")
            return

        if not self._is_file(source_markdown):
            self.show_error("Please select a valid source Markdown file.")
            return

        if not output_docx:
            self.show_error("Please select an output DOCX path.")
            return

        if not self._has_docx_extension(output_docx):
            self.show_error("Please select an output DOCX path ending in .docx.")
            return

        output_parent = Path(output_docx).expanduser().parent
        if output_parent and not output_parent.exists():
            self.show_error(f"Output folder does not exist:\n{output_parent}")
            return

        if os.path.exists(output_docx):
            should_overwrite = messagebox.askyesno(
                "Confirm Overwrite",
                "The output file already exists. Do you want to overwrite it?",
            )
            if not should_overwrite:
                self.log("Conversion cancelled. Output file already exists.")
                return

        metadata = {
            "{{CLIENT_NAME}}": self.client_name.get().strip(),
            "{{DOCUMENT_NAME}}": self.document_name.get().strip(),
            "{{AUTHOR}}": self.author.get().strip(),
        }
        self.log("Metadata values:")
        self.log(f"Client Name: {metadata['{{CLIENT_NAME}}'] or '(blank)'}")
        self.log(f"Document Name: {metadata['{{DOCUMENT_NAME}}'] or '(blank)'}")
        self.log(f"Author: {metadata['{{AUTHOR}}'] or '(blank)'}")

        try:
            self.generate_document(
                front_template_docx,
                reference_docx,
                source_markdown,
                output_docx,
                metadata,
            )
        except ProposalGenerationError as exc:
            self.show_error(str(exc))
        except Exception as exc:
            self.show_error(f"Unexpected error:\n{exc}")

    def generate_document(
        self,
        front_template_docx,
        reference_docx,
        source_markdown,
        output_docx,
        metadata,
    ):
        pandoc_path = find_pandoc()
        if not pandoc_path:
            raise ProposalGenerationError(
                "Pandoc is not installed in the local Python environment.\n"
                "Please run: .\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt"
            )

        version_line = check_pandoc_version(pandoc_path)
        self.log(f"Pandoc available: {version_line}")
        self.log(f"Pandoc path: {pandoc_path}")

        with tempfile.TemporaryDirectory(prefix="docx_proposal_") as temp_dir:
            body_docx = str(Path(temp_dir) / "body.docx")
            pandoc_command, pandoc_result = convert_markdown_to_docx(
                pandoc_path,
                source_markdown,
                reference_docx,
                body_docx,
            )

            self.log("Running Pandoc command:")
            self.log(self._format_command(pandoc_command))
            if pandoc_result.stdout:
                self.log("Pandoc stdout:")
                self.log(pandoc_result.stdout.strip())
            if pandoc_result.stderr:
                self.log("Pandoc stderr:")
                self.log(pandoc_result.stderr.strip())
            self.log(f"Converted Markdown body DOCX: {body_docx}")

            try:
                self.log("Composing final DOCX with Microsoft Word automation.")
                compose_with_word(
                    front_template_docx,
                    body_docx,
                    output_docx,
                    metadata,
                    self.log,
                )
                self.log(f"Success. DOCX generated at: {output_docx}")
                messagebox.showinfo("Success", "DOCX generated successfully.")
            except Exception as word_exc:
                self.log(f"Word automation unavailable or failed: {word_exc}")
                self.log("Falling back to python-docx composition.")
                compose_without_word(
                    front_template_docx,
                    body_docx,
                    output_docx,
                    metadata,
                    self.log,
                )
                self.log(f"Success with warning. DOCX generated at: {output_docx}")
                messagebox.showwarning(
                    "Generated With Warning",
                    "DOCX generated, but Microsoft Word was not available to update "
                    "the TOC, fields, and page numbers. Open the document in Word "
                    "and update fields if needed.",
                )

    def _is_file(self, path):
        return bool(path) and os.path.isfile(path)

    def _has_docx_extension(self, path):
        return Path(path).suffix.lower() == ".docx"

    def _has_front_template_extension(self, path):
        return Path(path).suffix.lower() in {".docx", ".dotx"}

    def _format_command(self, command):
        return subprocess.list2cmdline(command)

    def log(self, message):
        self.status_text.configure(state=tk.NORMAL)
        self.status_text.insert(tk.END, f"{message}\n")
        self.status_text.see(tk.END)
        self.status_text.configure(state=tk.DISABLED)

    def show_error(self, message):
        self.log(f"Error: {message}")
        messagebox.showerror("Error", message)


def main():
    app = DOCXProposalGenerator()
    app.mainloop()


if __name__ == "__main__":
    main()
